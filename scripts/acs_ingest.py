#!/usr/bin/env python3
"""
acs_ingest.py — ACS Planning Document Ingestion Script
=======================================================
Reads any existing planning document (BMAD, DevOps plan, product spec,
milestone doc, README, etc.) and automatically populates:

  - <acs-dir>/STATE.md
  - <acs-dir>/MUST_READ.md
  - <acs-dir>/MEMORY.md
  - <acs-dir>/PROTOCOL.md
  - CLAUDE.md / AGENT.md  (ACS trigger file)

Supports input formats:
  - Markdown (.md), Plain text (.txt), Word document (.docx)
  - Multiple files at once

Supported LLM providers:
  - anthropic  (default) — pip install anthropic
  - openai               — pip install openai
  - gemini               — pip install google-generativeai
  - ollama     (local)   — pip install ollama  +  ollama server running

Usage:
  python acs_ingest.py --input plan.md
  python acs_ingest.py --input plan.md --provider openai --model gpt-4o
  python acs_ingest.py --input plan.md --provider gemini --model gemini-1.5-pro
  python acs_ingest.py --input plan.md --provider ollama --model llama3
  python acs_ingest.py --input plan.md --provider auto        # detect from env
  python acs_ingest.py --input plan.md --acs-dir .acs --agent-file AGENT.md

Auto-detection priority (--provider auto):
  ANTHROPIC_API_KEY → openai: OPENAI_API_KEY → gemini: GEMINI_API_KEY → ollama
"""

import argparse
import json
import os
import re
import sys
from abc import ABC, abstractmethod
from datetime import datetime, timezone
from pathlib import Path

# ── optional imports ──────────────────────────────────────────────────────────
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

NOW = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")

DEFAULT_MODELS = {
    "anthropic": "claude-sonnet-4-20250514",
    "openai":    "gpt-4o",
    "gemini":    "gemini-1.5-pro",
    "ollama":    "llama3",
}

EXTRACTION_PROMPT = """You are reading a software project planning document.
Extract structured information to populate an ACS project tracking setup.

Return ONLY valid JSON — no explanation, no markdown fences, no preamble:

{{
  "project_name": "string",
  "project_description": "string — 1-2 sentences",
  "tech_stack": ["list", "of", "technologies"],
  "phases": [
    {{"phase_number": 1, "phase_name": "string", "phase_description": "string"}}
  ],
  "milestones": [
    {{
      "id": "M-001",
      "phase": 1,
      "name": "string",
      "description": "string",
      "category": "Infrastructure|DataPipeline|FeatureEngineering|Model|API|Testing|Deployment|Other",
      "size": "SMALL|MEDIUM|LARGE",
      "dependencies": []
    }}
  ],
  "components": [
    {{"name": "string", "description": "string", "category": "string"}}
  ],
  "architectural_decisions": [
    {{"title": "string", "decision": "string", "rationale": "string"}}
  ],
  "external_dependencies": [
    {{"name": "string", "purpose": "string", "version_or_notes": "string"}}
  ],
  "key_constraints": ["string"],
  "first_session_tasks": ["string — first 3-5 things to do, in order"],
  "branch_strategy": "string",
  "environment_notes": "string"
}}

Rules:
- Use "Unknown" for anything not in the document
- Milestone IDs: M-001, M-002, ...
- Size: SMALL=single func/config, MEDIUM=module+tests, LARGE=multi-component
- Infer phases from logical groupings if not explicit
- Extract ALL milestones, even implicit ones
- Do not invent information not present in the document

Document:
---
{document_content}
---"""


# ══════════════════════════════════════════════════════════════════════════════
# LLM CLIENT ABSTRACTION
# ══════════════════════════════════════════════════════════════════════════════

class LLMClient(ABC):
    def __init__(self, api_key: str | None, model: str):
        self.api_key = api_key
        self.model = model

    @abstractmethod
    def complete(self, prompt: str, max_tokens: int = 4000) -> str: ...

    @classmethod
    @abstractmethod
    def provider_name(cls) -> str: ...

    @classmethod
    @abstractmethod
    def env_var(cls) -> str | None: ...


class AnthropicClient(LLMClient):
    @classmethod
    def provider_name(cls): return "anthropic"
    @classmethod
    def env_var(cls): return "ANTHROPIC_API_KEY"

    def __init__(self, api_key, model):
        super().__init__(api_key, model)
        try:
            import anthropic as _a
        except ImportError:
            _die("anthropic", "pip install anthropic")
        self._client = _a.Anthropic(api_key=self.api_key)

    def complete(self, prompt, max_tokens=4000):
        r = self._client.messages.create(
            model=self.model, max_tokens=max_tokens,
            messages=[{"role": "user", "content": prompt}])
        return r.content[0].text.strip()


class OpenAIClient(LLMClient):
    @classmethod
    def provider_name(cls): return "openai"
    @classmethod
    def env_var(cls): return "OPENAI_API_KEY"

    def __init__(self, api_key, model):
        super().__init__(api_key, model)
        try:
            import openai as _o
        except ImportError:
            _die("openai", "pip install openai")
        self._client = _o.OpenAI(api_key=self.api_key)

    def complete(self, prompt, max_tokens=4000):
        r = self._client.chat.completions.create(
            model=self.model, max_tokens=max_tokens,
            messages=[{"role": "user", "content": prompt}])
        return r.choices[0].message.content.strip()


class GeminiClient(LLMClient):
    @classmethod
    def provider_name(cls): return "gemini"
    @classmethod
    def env_var(cls): return "GEMINI_API_KEY"

    def __init__(self, api_key, model):
        super().__init__(api_key, model)
        try:
            import google.generativeai as _g
        except ImportError:
            _die("google-generativeai", "pip install google-generativeai")
        _g.configure(api_key=self.api_key)
        self._model = _g.GenerativeModel(model)

    def complete(self, prompt, max_tokens=4000):
        r = self._model.generate_content(
            prompt, generation_config={"max_output_tokens": max_tokens})
        return r.text.strip()


class OllamaClient(LLMClient):
    """Local Ollama — no API key required."""
    @classmethod
    def provider_name(cls): return "ollama"
    @classmethod
    def env_var(cls): return None

    def __init__(self, api_key, model):
        super().__init__(api_key, model)
        try:
            import ollama as _ol
            self._ol = _ol
        except ImportError:
            _die("ollama", "pip install ollama  (and ensure ollama server is running)")

    def complete(self, prompt, max_tokens=4000):
        r = self._ol.chat(model=self.model,
                          messages=[{"role": "user", "content": prompt}])
        return r["message"]["content"].strip()


def _die(pkg, install_cmd):
    print(f"ERROR: {pkg} package not installed.\nRun:   {install_cmd}")
    sys.exit(1)


PROVIDERS: dict[str, type[LLMClient]] = {
    "anthropic": AnthropicClient,
    "openai":    OpenAIClient,
    "gemini":    GeminiClient,
    "ollama":    OllamaClient,
}


def resolve_provider(provider: str, api_key_override: str | None) -> tuple[str, str | None]:
    """Returns (resolved_provider_name, api_key)."""
    if provider == "auto":
        for name in ["anthropic", "openai", "gemini"]:
            env = PROVIDERS[name].env_var()
            if env and os.environ.get(env):
                print(f"  Auto-detected provider: {name} (from {env})")
                provider = name
                break
        else:
            print("  No API key env vars found — falling back to ollama (local)")
            provider = "ollama"

    if provider not in PROVIDERS:
        print(f"ERROR: Unknown provider '{provider}'. Choose: {', '.join(PROVIDERS)}")
        sys.exit(1)

    env_var = PROVIDERS[provider].env_var()
    api_key = api_key_override or (os.environ.get(env_var) if env_var else None)

    # Try .env file
    if api_key is None and env_var:
        env_file = Path(".env")
        if env_file.exists():
            for line in env_file.read_text().splitlines():
                if line.startswith(f"{env_var}="):
                    api_key = line.split("=", 1)[1].strip().strip('"\'')
                    break

    if api_key is None and env_var:
        print(f"ERROR: {env_var} not set.\nSet it as an env var, pass --api-key, or add to .env")
        sys.exit(1)

    return provider, api_key


# ══════════════════════════════════════════════════════════════════════════════
# DOCUMENT READERS
# ══════════════════════════════════════════════════════════════════════════════

def read_input_file(path: Path) -> str:
    if path.suffix.lower() == ".docx":
        if not DOCX_AVAILABLE:
            print(f"ERROR: python-docx not installed.\nRun: pip install python-docx")
            sys.exit(1)
        doc = DocxDocument(str(path))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                s = para.style.name
                prefix = "# " if "Heading 1" in s else "## " if "Heading 2" in s else "### " if "Heading 3" in s else ""
                parts.append(f"{prefix}{para.text}")
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(c.text.strip() for c in row.cells))
        return "\n".join(parts)
    try:
        return path.read_text(encoding="utf-8")
    except Exception:
        print(f"WARNING: Cannot read {path} — skipping.")
        return ""


# ══════════════════════════════════════════════════════════════════════════════
# EXTRACTION
# ══════════════════════════════════════════════════════════════════════════════

def extract_structure(document_content: str, client: LLMClient) -> dict:
    print(f"  Analysing with {client.provider_name()} / {client.model} ...")
    raw = client.complete(
        EXTRACTION_PROMPT.format(document_content=document_content[:50000]))
    raw = re.sub(r'^```(?:json)?\s*', '', raw)
    raw = re.sub(r'\s*```$', '', raw)
    try:
        return json.loads(raw)
    except json.JSONDecodeError as e:
        print(f"ERROR: LLM returned invalid JSON: {e}")
        print("Raw response (first 500 chars):", raw[:500])
        sys.exit(1)


# ══════════════════════════════════════════════════════════════════════════════
# DOCUMENT GENERATORS
# ══════════════════════════════════════════════════════════════════════════════

def generate_state_md(data: dict) -> str:
    milestones  = data.get("milestones", [])
    phases      = data.get("phases", [])
    phase_names = {p["phase_number"]: p["phase_name"] for p in phases}

    by_phase: dict = {}
    for m in milestones:
        by_phase.setdefault(m.get("phase", 1), {}).setdefault(
            m.get("category", "Other"), []).append(m)

    lines = [
        "# PROJECT STATE — VERIFIED COMPLETIONS ONLY",
        "**Rule**: Nothing is added here without a verified commit hash.",
        f"**Last verification run**: {NOW}",
        f"**Project**: {data.get('project_name', 'Unknown')}",
        f"**Auto-generated**: {NOW}",
        "", "---", "",
        "## LEGEND",
        "| Status | Meaning |", "|--------|---------|",
        "| `VERIFIED` | Confirmed — commit hash and test evidence recorded |",
        "| `IN_PROGRESS` | Currently being built this session |",
        "| `PARTIAL` | Started but not complete — see CHECKPOINT.md |",
        "| `NOT_STARTED` | On the roadmap but not yet begun |",
        "| `BLOCKED` | Cannot proceed — blocker in MUST_READ.md |",
        "", "---", "",
    ]

    for phase_num in sorted(by_phase):
        lines.append(f"## Phase {phase_num}: {phase_names.get(phase_num, f'Phase {phase_num}')}")
        lines.append("")
        for cat, items in by_phase[phase_num].items():
            lines += [f"### {cat}", "",
                      "| ID | Milestone | Size | Status | Commit | Verified |",
                      "|-----|-----------|------|--------|--------|----------|"]
            for m in items:
                lines.append(
                    f"| {m.get('id','M-???')} | {m.get('name','Unknown')[:50]} "
                    f"| {m.get('size','MEDIUM')} | NOT_STARTED | — | — |")
                if m.get("description"):
                    lines.append(f"| | *{m['description'][:80]}* | | | | |")
            lines.append("")

    ext_deps = data.get("external_dependencies", [])
    lines += [
        "---", "", "## Infrastructure & Environment", "",
        "| Component | Status | Commit | Verified Timestamp |",
        "|-----------|--------|--------|--------------------|",
        f"| Git repository | VERIFIED | (initial) | {NOW} |",
        f"| ACS Protocol files | VERIFIED | (initial) | {NOW} |",
        "| Python environment | NOT_STARTED | — | — |",
    ]
    if ext_deps:
        lines += ["", "---", "", "## External Dependencies", "",
                  "| Dependency | Purpose | Version / Notes | Status |",
                  "|------------|---------|-----------------|--------|"]
        for d in ext_deps:
            lines.append(
                f"| {d.get('name','')} | {d.get('purpose','')} "
                f"| {d.get('version_or_notes','—')} | NOT_STARTED |")
    lines += [
        "", "---", "", "## Test Suites", "",
        "| Test Suite | Pass | Fail | Last Run | Commit |",
        "|------------|------|------|----------|--------|",
        "| (no tests yet) | — | — | — | — |",
        "", "---", "", "## Git State",
        "- Current branch: main",
        "- Last push confirmed: (not yet pushed)",
        "- Remote sync status: NOT_CONFIGURED", "",
    ]
    return "\n".join(lines)


def generate_must_read_md(data: dict) -> str:
    first_tasks = data.get("first_session_tasks", [])
    phases = data.get("phases", [])
    fp = phases[0] if phases else {"phase_number": 1, "phase_name": "Setup"}
    task_lines = [f"{i}. ATU-00{i}: {t} — SMALL"
                  for i, t in enumerate(first_tasks[:5], 1)] or [
        "1. ATU-001: Run verify_state.py and confirm clean baseline — SMALL",
        "2. ATU-002: Set up Python environment and install dependencies — SMALL",
        "3. ATU-003: Configure git remote and confirm push works — SMALL",
    ]
    lines = [
        "# MUST READ — SESSION STARTUP BRIEF",
        f"**Last updated**: {NOW}",
        "**Updated by**: ACS Ingestion (auto-generated)",
        "**Emergency flag**: NONE", "",
        "## ⚡ IMMEDIATE ACTIONS",
        "1. Run: `python .claude/scripts/verify_state.py`",
        "2. If verification fails: resolve before ANY new work",
        "3. Read CHECKPOINT.md if Emergency flag is set", "",
        "---", "",
        f"## Project: {data.get('project_name', 'Unknown')}",
        data.get("project_description", ""), "",
        "## Current Phase",
        f"**Phase**: {fp.get('phase_number',1)} — {fp.get('phase_name','Setup')}",
        "**Blocking Issues**: NONE", "",
        "## Last Verified Completion",
        "**ATU**: ATU-000 — ACS Protocol Initialised",
        f"**Timestamp**: {NOW}", "",
        "## THIS Session's Tasks (in order)",
        "\n".join(task_lines), "",
        "## Architecture Summary",
    ]
    if data.get("tech_stack"):
        lines.append(f"- **Tech stack**: {', '.join(data['tech_stack'])}")
    env = data.get("environment_notes", "")
    if env and env != "Unknown":
        lines.append(f"- **Environment**: {env}")
    branch = data.get("branch_strategy", "")
    if branch and branch != "Unknown":
        lines.append(f"- **Branch strategy**: {branch}")
    for c in data.get("key_constraints", [])[:3]:
        lines.append(f"  - {c}")
    lines += [
        "", "## Do NOT Do This Session",
        "- Write any application code before ACS baseline is confirmed",
        "- Mark anything VERIFIED in STATE.md without a commit hash",
        "- Start a LARGE task without confirming sufficient session capacity",
    ]
    return "\n".join(lines)


def generate_memory_md(data: dict) -> str:
    lines = [
        "# PROJECT MEMORY — PERSISTENT CONTEXT",
        f"**Project**: {data.get('project_name', 'Unknown')}",
        f"**Initialised**: {NOW}", "", "---", "", "## Architectural Decisions", "",
        f"### {NOW}: ACS Protocol Adopted",
        "**Decision**: Using the Absolute Continuity System (ACS) for all session management.",
        "**Rationale**: Enforces verification gates before marking anything complete.",
        "**Impact**: Every task requires verify command + git log confirmation.", "",
    ]
    for dec in data.get("architectural_decisions", []):
        lines += [
            f"### {NOW}: {dec.get('title','Untitled')}",
            f"**Decision**: {dec.get('decision','')}",
            f"**Rationale**: {dec.get('rationale','')}",
            "**Alternatives considered**: (document as they arise)", "",
        ]
    ext_deps = data.get("external_dependencies", [])
    lines += [
        "---", "", "## Problems Encountered and Resolved",
        "(None yet — add as they occur)", "", "---", "",
        "## External Dependencies", "",
        "| Dependency | Version | Purpose | First encountered |",
        "|------------|---------|---------|-------------------|",
    ]
    for d in ext_deps:
        lines.append(
            f"| {d.get('name','')} | {d.get('version_or_notes','—')} "
            f"| {d.get('purpose','')} | {NOW} |")
    if not ext_deps:
        lines.append("| (none yet) | — | — | — |")
    lines += [
        "", "---", "", "## Tech Stack Reference", "",
    ]
    lines += [f"- {t}" for t in data.get("tech_stack", [])] or ["(to be documented)"]
    return "\n".join(lines)


def generate_agent_md(data: dict, acs_dir: str, agent_filename: str) -> str:
    project_name = data.get("project_name", "Unknown Project")
    project_desc = data.get("project_description", "")
    tech_stack   = data.get("tech_stack", [])
    stack_str    = ", ".join(tech_stack[:5]) or "see MEMORY.md"
    phases       = data.get("phases", [])
    fp           = phases[0] if phases else {"phase_number": 1, "phase_name": "Setup"}

    return f"""# {agent_filename} — ACS PROTOCOL TRIGGER
# Keep this file under 100 lines. All detail lives in {acs_dir}/

## ⚡ MANDATORY STARTUP — Do this before ANY other action

### Step 1: Run verification
```bash
python {acs_dir}/scripts/verify_state.py
```

### Step 2: Read session brief
Read `{acs_dir}/MUST_READ.md` in full.

### Step 3: First response must confirm ALL of:
- [ ] Verification result: exit 0 (clean) or exit 1 (discrepancies found)
- [ ] Current git HEAD: output of `git log --oneline -1`
- [ ] Last verified ATU: from STATE.md
- [ ] CHECKPOINT.md status: clean / recovery needed
- [ ] Planned tasks this session: from MUST_READ.md

**Do not write any code until all five confirmations are in your response.**

---

## Project: {project_name}
{project_desc}

## Phase {fp.get('phase_number', 1)} — {fp.get('phase_name', 'Setup')}

## Tech Stack
{stack_str}

## ACS Document Map
| Need | File |
|------|------|
| What to do this session | `{acs_dir}/MUST_READ.md` |
| What is verified complete | `{acs_dir}/STATE.md` |
| Architectural decisions | `{acs_dir}/MEMORY.md` |
| Live session progress | `{acs_dir}/CHECKPOINT.md` |
| Protocol quick reference | `{acs_dir}/PROTOCOL.md` |

## Never Do Without Permission
- Merge to main without all tests passing
- Mark STATE.md VERIFIED without a commit hash
- Start a LARGE task without a credit horizon check
- End a session without updating MUST_READ.md for next session
"""


def generate_protocol_md() -> str:
    return """# ACS PROTOCOL — QUICK REFERENCE

## Every Session Starts With:
1. Read MUST_READ.md
2. Read STATE.md
3. Read CHECKPOINT.md (if it exists)
4. Run: python .claude/scripts/verify_state.py
5. Resolve ALL discrepancies before writing any new code

## A Task (ATU) is DONE When:
- [ ] Verify command run and output matches expected
- [ ] git commit run with standard message
- [ ] git log --oneline -1 confirms commit landed
- [ ] STATE.md updated with commit hash and timestamp
- [ ] CHECKPOINT.md updated to show task COMPLETE

## Never Say "Complete" Without:
Running verify command AND showing actual output
AND git log --oneline -1 AND showing that output.

## Credit Check (before MEDIUM or LARGE tasks):
"Task size: [SIZE]. Estimated: [X min]. Capacity: [assessment]. Safe: YES/NO"

## Session End (Planned):
1. Complete current ATU fully
2. Update CHECKPOINT.md → COMPLETED
3. Update STATE.md with verified completions
4. Add learnings to MEMORY.md
5. Write next session's MUST_READ.md
6. git push && git status (clean working tree)
7. Say: "SESSION TERMINATION COMPLETE. Next: ATU-[X]"

## Emergency End:
1. git add -A && git commit -m "WIP: EMERGENCY [state]" && git push
2. Append emergency block to CHECKPOINT.md
3. Add WARNING flag to top of MUST_READ.md

## The Ten Golden Rules:
1.  Verify before trusting
2.  Never claim complete — show verification output
3.  Commit hash or it didn't happen
4.  STATE.md = reality not intention
5.  CHECKPOINT.md is live — update after every ATU
6.  Check credit before large tasks
7.  80% rule — finish cleanly
8.  Emergency commit — always better than undocumented state
9.  Partial is documented — never leave undocumented partial state
10. MUST_READ.md written before session ends
"""


# ══════════════════════════════════════════════════════════════════════════════
# FILE WRITING
# ══════════════════════════════════════════════════════════════════════════════

def write_file_safe(path: Path, content: str, dry_run: bool, force: bool) -> bool:
    if path.exists() and not force:
        if path.read_text(encoding="utf-8") == content:
            print(f"  UNCHANGED  {path}")
            return False
        print(f"  MODIFIED   {path}  (already exists)")
        if not dry_run:
            if input(f"    Overwrite {path}? [y/N]: ").strip().lower() != "y":
                print("    Skipped.")
                return False
    if dry_run:
        print(f"  DRY-RUN    Would write {path} ({len(content)} chars)")
        return False
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(content, encoding="utf-8")
    print(f"  WRITTEN    {path}")
    return True


# ══════════════════════════════════════════════════════════════════════════════
# MAIN
# ══════════════════════════════════════════════════════════════════════════════

def main():
    parser = argparse.ArgumentParser(
        description="ACS Planning Document Ingestion — multi-provider LLM support.")
    parser.add_argument("--input", "-i", nargs="+", required=True, metavar="FILE")
    parser.add_argument("--provider", "-p", default="auto",
                        choices=list(PROVIDERS) + ["auto"],
                        help="LLM provider. Default: auto (detect from env vars)")
    parser.add_argument("--model", "-m", default=None,
                        help=f"Model override. Defaults: {DEFAULT_MODELS}")
    parser.add_argument("--api-key", default=None,
                        help="API key override (default: read from env var)")
    parser.add_argument("--project-name", "-n", default=None,
                        help="Override extracted project name")
    parser.add_argument("--acs-dir", default=".claude",
                        help="ACS directory (default: .claude). Use .acs for provider-neutral setup.")
    parser.add_argument("--agent-file", default="CLAUDE.md",
                        help="Agent trigger filename (default: CLAUDE.md). Use AGENT.md for neutral setup.")
    parser.add_argument("--dry-run", action="store_true")
    parser.add_argument("--force", "-f", action="store_true")
    parser.add_argument("--output-dir", default=".")
    parser.add_argument("--skip-agent-file", action="store_true")
    parser.add_argument("--json-only", action="store_true",
                        help="Print extracted JSON only, do not write files")
    args = parser.parse_args()

    print("\nACS Document Ingestion")
    print("=" * 60)

    resolved_provider, api_key = resolve_provider(args.provider, args.api_key)
    model = args.model or DEFAULT_MODELS[resolved_provider]
    client = PROVIDERS[resolved_provider](api_key=api_key, model=model)
    print(f"  Provider:  {resolved_provider}")
    print(f"  Model:     {client.model}")

    print("\nReading input files...")
    combined = []
    for f in args.input:
        p = Path(f)
        if not p.exists():
            print(f"  ERROR: Not found: {p}")
            sys.exit(1)
        print(f"  Reading:   {p} ({p.stat().st_size // 1024}KB)")
        text = read_input_file(p)
        if text:
            combined.append(f"\n\n--- SOURCE: {p.name} ---\n\n{text}")

    if not combined:
        print("ERROR: No readable content found.")
        sys.exit(1)

    full_doc = "\n".join(combined)
    print(f"  Total:     {len(full_doc):,} chars from {len(combined)} file(s)")

    print("\nExtracting project structure...")
    data = extract_structure(full_doc, client)
    if args.project_name:
        data["project_name"] = args.project_name

    print(f"  Project:         {data.get('project_name', 'Unknown')}")
    print(f"  Phases:          {len(data.get('phases', []))}")
    print(f"  Milestones:      {len(data.get('milestones', []))}")
    print(f"  Arch decisions:  {len(data.get('architectural_decisions', []))}")
    print(f"  External deps:   {len(data.get('external_dependencies', []))}")

    if args.json_only:
        print(json.dumps(data, indent=2))
        return

    output_root = Path(args.output_dir)
    acs_dir     = output_root / args.acs_dir

    print(f"\nGenerating ACS documents → {output_root.resolve()}")
    print(f"  ACS dir:   {args.acs_dir}/")
    print(f"  Agent file: {args.agent_file}")
    if args.dry_run:
        print("  (DRY RUN)")
    print()

    written = []
    for fname, gen in [
        ("STATE.md",    generate_state_md(data)),
        ("MUST_READ.md", generate_must_read_md(data)),
        ("MEMORY.md",   generate_memory_md(data)),
        ("PROTOCOL.md", generate_protocol_md()),
    ]:
        if write_file_safe(acs_dir / fname, gen, args.dry_run, args.force):
            written.append(fname)

    if not args.skip_agent_file:
        agent_content = generate_agent_md(data, args.acs_dir, args.agent_file)
        if write_file_safe(output_root / args.agent_file, agent_content,
                           args.dry_run, args.force):
            written.append(args.agent_file)

    if not args.dry_run:
        (acs_dir / "scripts").mkdir(parents=True, exist_ok=True)
        jp = acs_dir / "ingestion_result.json"
        jp.write_text(json.dumps(data, indent=2))
        print(f"  WRITTEN    {jp}")
        written.append("ingestion_result.json")

    print()
    print("=" * 60)
    if args.dry_run:
        print("DRY RUN COMPLETE — no files written")
    else:
        print(f"INGESTION COMPLETE — {len(written)} files written")
        print(f"\n  {len(data.get('milestones',[]))} milestones added to STATE.md (all NOT_STARTED)")
        print(f"\nNEXT STEPS:")
        print(f"  1. Review {args.acs_dir}/STATE.md")
        print(f"  2. Review {args.acs_dir}/MUST_READ.md")
        print(f"  3. python {args.acs_dir}/scripts/verify_state.py")
        print(f"  4. git add {args.acs_dir}/ {args.agent_file} && git commit")
    print("=" * 60)


if __name__ == "__main__":
    main()
